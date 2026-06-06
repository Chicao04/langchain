import os
import sys
import glob
import pdfplumber
import pandas as pd
from docx import Document

# Thiết lập mã hóa hiển thị cho Terminal Windows tránh lỗi font chữ tiếng Việt
if sys.platform.startswith('win'):
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')

# =====================================================================
# KHỐI CẤU HÌNH HỆ THỐNG (BẠN CHỈ CẦN CHỈNH SỬA TẠI ĐÂY)
# =====================================================================
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(CURRENT_DIR) 

CONFIG = {
    # ĐƯỜNG DẪN THƯ MỤC ĐẦU VÀO: Nơi chứa nhiều file PDF và DOCX của bạn
    "INPUT_DIR_PATH": os.path.join(PROJECT_ROOT, "input_data"),
    
    # ĐƯỜNG DẪN THƯ MỤC ĐẦU RA: Nơi lưu kết quả trích xuất
    "OUTPUT_DIR_PATH": os.path.join(PROJECT_ROOT, "output_data", "extract_transform_data_output")
}

# Tự động kiểm tra và khởi tạo cây thư mục đầu ra nếu chưa tồn tại
os.makedirs(CONFIG["OUTPUT_DIR_PATH"], exist_ok=True)


# =====================================================================
# 1. LỚP XỬ LÝ TRÍCH XUẤT DỮ LIỆU (DATA EXTRACTION)
# =====================================================================
class DocumentTableExtractor:
    """Lớp chuyên dụng chịu trách nhiệm trích xuất cấu trúc dữ liệu bảng và văn bản thô."""
    
    @staticmethod
    def extract_from_docx(file_path: str) -> tuple[list[pd.DataFrame], str]:
        """Trích xuất bảng và văn bản thô từ file Word, loại bỏ hoàn toàn text bảng khỏi văn bản thô."""
        doc = Document(file_path)
        danh_sach_bang = []
        
        text_trong_cac_bang = set()
        for idx, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                text_in_row = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    text_in_row.append(cell_text)
                    if cell_text:
                        text_trong_cac_bang.add(cell_text)
                data.append(text_in_row)
            
            if data and len(data) > 1:
                df = pd.DataFrame(data[1:], columns=data[0])
                danh_sach_bang.append(df)
        
        van_ban_list = []
        for p in doc.paragraphs:
            text_p = p.text.strip()
            if text_p and text_p not in text_trong_cac_bang:
                van_ban_list.append(text_p)
                
        van_ban_thong_thuong = "\n".join(van_ban_list)
        return danh_sach_bang, van_ban_thong_thuong

    @staticmethod
    def extract_from_pdf(file_path: str) -> tuple[list[pd.DataFrame], str]:
        """Trích xuất bảng và văn bản thô từ file PDF, xóa sạch vùng chữ thuộc bảng trước khi lấy text thô."""
        danh_sach_bang = []
        toan_bo_text = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                tables_found = page.find_tables()
                
                def xoa_vung_chua_bang(obj):
                    if obj.get("object_type") == "char":
                        for table in tables_found:
                            x0, y0, x1, y1 = table.bbox
                            if x0 <= obj["x0"] <= x1 and y0 <= obj["top"] <= y1:
                                return False
                    return True

                trang_da_xoa_bang = page.filter(xoa_vung_chua_bang)
                text_ngoai_bang = trang_da_xoa_bang.extract_text()
                if text_ngoai_bang and text_ngoai_bang.strip():
                    toan_bo_text.append(f"--- TEXT TRANG {page_idx + 1} ---\n{text_ngoai_bang.strip()}")
                
                tables = page.extract_tables()
                for t_idx, table in enumerate(tables):
                    if table and len(table) > 1:
                        cleaned_table = [
                            [str(cell).strip() if cell is not None else "" for cell in row] 
                            for row in table
                        ]
                        df = pd.DataFrame(cleaned_table[1:], columns=cleaned_table[0])
                        danh_sach_bang.append(df)
                        
        return danh_sach_bang, "\n\n".join(toan_bo_text)


# =====================================================================
# 2. LỚP LƯU TRỮ DỮ LIỆU ĐẦU RA (DATA STORAGE SERVICE)
# =====================================================================
class DataStorageService:
    """Lớp lưu trữ chịu trách nhiệm ghi file kết quả động theo tên từng file đầu vào."""
    
    @staticmethod
    def save_tables_to_markdown(danh_sach_bang: list[pd.DataFrame], base_filename: str):
        """Lưu toàn bộ các bảng thu được thành định dạng lưới Markdown vào thư mục đích."""
        output_filename = f"{base_filename}_bang_raw.txt"
        full_output_path = os.path.join(CONFIG["OUTPUT_DIR_PATH"], output_filename)
        
        if not danh_sach_bang:
            with open(full_output_path, "w", encoding="utf-8") as f:
                f.write("Không tìm thấy dữ liệu dạng bảng nào trong tài liệu.")
            return

        with open(full_output_path, "w", encoding="utf-8") as f:
            f.write("==================================================\n")
            f.write(f"   KẾT QUẢ TRÍCH XUẤT DỮ LIỆU BẢNG: {base_filename}\n")
            f.write("==================================================\n\n")
            
            for idx, df in enumerate(danh_sach_bang):
                f.write(f"--- BẢNG SỐ {idx + 1} ---\n")
                f.write(f"Kích thước: {df.shape[0]} hàng x {df.shape[1]} cột\n")
                f.write(f"Cột: {list(df.columns)}\n\n")
                f.write(df.to_markdown(index=False))
                f.write("\n\n" + "-" * 60 + "\n\n")
                
        print(f"   -> Đã xuất cấu trúc bảng vào: '{output_filename}'")

    @staticmethod
    def save_text_to_file(van_ban_thong_thuong: str, base_filename: str):
        """Lưu trữ văn bản chữ thường không thuộc bảng vào thư mục đích."""
        output_filename = f"{base_filename}_text_raw.txt"
        full_output_path = os.path.join(CONFIG["OUTPUT_DIR_PATH"], output_filename)
        
        with open(full_output_path, "w", encoding="utf-8") as f:
            f.write("==================================================\n")
            f.write(f"   KẾT QUẢ TRÍCH XUẤT VĂN BẢN (ĐÃ LỌC BẢNG): {base_filename}\n")
            f.write("==================================================\n\n")
            f.write(van_ban_thong_thuong)
        print(f"   -> Đã xuất cấu trúc chữ thô vào: '{output_filename}'")


# =====================================================================
# 3. HÀM ĐIỀU PHỐI PIPELINE CHÍNH (XỬ LÝ HÀNG LOẠT)
# =====================================================================
def execute_batch_pipeline():
    """Hàm quét thư mục, nhận diện toàn bộ file và điều phối bóc tách hàng loạt."""
    input_dir = CONFIG["INPUT_DIR_PATH"]
    
    if not os.path.exists(input_dir):
        print(f"[Lỗi] Thư mục đầu vào không tồn tại: {input_dir}")
        return
        
    # Lấy tất cả các file có đuôi .pdf và .docx trong thư mục đầu vào
    valid_extensions = ["*.pdf", "*.docx"]
    all_files = []
    for ext in valid_extensions:
        all_files.extend(glob.glob(os.path.join(input_dir, ext)))
        
    if not all_files:
        print(f"[Thông báo] Không tìm thấy file .pdf hoặc .docx nào trong thư mục: {input_dir}")
        return
        
    print(f"[*] Tìm thấy tất cả {len(all_files)} file tài liệu cần xử lý.")
    print("="*60)

    # Vòng lặp duyệt qua từng file tài liệu
    for file_idx, file_path in enumerate(all_files, start=1):
        filename_with_ext = os.path.basename(file_path)
        base_filename, file_extension = os.path.splitext(filename_with_ext)
        
        print(f"\n[{file_idx}/{len(all_files)}] Đang tiến hành bóc tách: {filename_with_ext}")
        
        danh_sach_bang = []
        van_ban_thong_thuong = ""
        
        # Rẽ nhánh bóc tách theo định dạng file cụ thể
        if file_extension.lower() == ".docx":
            danh_sach_bang, van_ban_thong_thuong = DocumentTableExtractor.extract_from_docx(file_path)
        elif file_extension.lower() == ".pdf":
            danh_sach_bang, van_ban_thong_thuong = DocumentTableExtractor.extract_from_pdf(file_path)
        else:
            continue

        # Lưu kết quả đầu ra theo tên động của file gốc để không bị đè dữ liệu
        DataStorageService.save_tables_to_markdown(danh_sach_bang, base_filename)
        DataStorageService.save_text_to_file(van_ban_thong_thuong, base_filename)
        
    print("\n" + "="*60 + "\n[HOÀN TẤT] Đã xử lý thành công toàn bộ các file tài liệu thô.\n" + "="*60)


if __name__ == "__main__":
    execute_batch_pipeline()